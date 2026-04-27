from __future__ import annotations

import json
import os
import shutil
import uuid
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path

import requests
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from docx import Document as DocxDocument
from pypdf import PdfReader

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate

from langchain_chroma import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings


# =========================
# ENV & PATHS
# =========================
BASE_DIR = Path(__file__).resolve().parent
# load_dotenv(BASE_DIR / ".env")
load_dotenv()

DATA_DIR      = BASE_DIR / "data"
UPLOAD_DIR    = DATA_DIR / "uploads"
VECTOR_DIR    = DATA_DIR / "chroma"
STATE_PATH    = DATA_DIR / "documents.json"
HISTORY_PATH  = DATA_DIR / "conversations.json"

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# Chunking — larger chunks keep more context together
CHUNK_SIZE    = 1200
CHUNK_OVERLAP = 300

# Retrieval
RETRIEVE_K      = 12   # fetch this many candidates from vector DB
FINAL_K         = 6    # keep top N after reranking
MIN_SCORE       = 0.3  # discard chunks below this similarity score (0–1)


# =========================
# INIT HELPERS
# =========================
def ensure_dirs():
    DATA_DIR.mkdir(exist_ok=True)
    UPLOAD_DIR.mkdir(exist_ok=True)
    VECTOR_DIR.mkdir(exist_ok=True)
    if not STATE_PATH.exists():
        STATE_PATH.write_text("[]")
    if not HISTORY_PATH.exists():
        HISTORY_PATH.write_text("{}")

def now():
    return datetime.now(timezone.utc).isoformat()

def load_state():
    ensure_dirs()
    return json.loads(STATE_PATH.read_text())

def save_state(data):
    STATE_PATH.write_text(json.dumps(data, indent=2))

def load_history() -> dict:
    ensure_dirs()
    return json.loads(HISTORY_PATH.read_text())

def save_history(data: dict):
    HISTORY_PATH.write_text(json.dumps(data, indent=2))


# =========================
# GEMINI MODEL DETECTION
# =========================
def list_all_gemini_models(api_key: str) -> list[dict]:
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
    try:
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        return resp.json().get("models", [])
    except Exception as e:
        print(f"[AutoDetect] Could not list models: {e}")
        return []

EMBEDDING_CANDIDATES = [
    "models/gemini-embedding-001",
    "models/gemini-embedding-2-preview",
    "models/text-embedding-004",
    "models/embedding-001",
]

def detect_embedding_backend(api_key: str, all_models: list[dict]):
    supported_names = [
        m["name"] for m in all_models
        if "embedContent" in m.get("supportedGenerationMethods", [])
    ]
    print(f"[AutoDetect] Embedding models available: {supported_names}")

    for candidate in EMBEDDING_CANDIDATES:
        if candidate in supported_names:
            print(f"[AutoDetect] ✅ Using Gemini embedding: {candidate}")
            emb = GoogleGenerativeAIEmbeddings(
                model=candidate,
                task_type="retrieval_document",
                google_api_key=api_key,
            )
            return emb, f"Gemini ({candidate})"

    print("[AutoDetect] ⚠️  Falling back to local embeddings...")
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
        emb = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
        print("[AutoDetect] ✅ Local embeddings ready.")
        return emb, "Local (all-MiniLM-L6-v2)"
    except ImportError:
        raise RuntimeError("Run: pip install sentence-transformers langchain-huggingface")

LLM_CANDIDATES = [
    "models/gemini-2.5-flash",
    "models/gemini-2.0-flash",
    "models/gemini-2.0-flash-lite",
    "models/gemini-1.5-flash",
    "models/gemini-1.5-pro",
    "models/gemini-pro",
]

def detect_llm(api_key: str, all_models: list[dict]) -> tuple[str, str]:
    supported_names = [
        m["name"] for m in all_models
        if "generateContent" in m.get("supportedGenerationMethods", [])
    ]
    print(f"[AutoDetect] LLM models available: {supported_names}")
    for candidate in LLM_CANDIDATES:
        if candidate in supported_names:
            print(f"[AutoDetect] ✅ Using Gemini LLM: {candidate}")
            return candidate, candidate
    if supported_names:
        chosen = supported_names[0]
        print(f"[AutoDetect] ✅ Using first available LLM: {chosen}")
        return chosen, chosen
    raise RuntimeError("No Gemini LLM found for your API key!")


# =========================
# STARTUP
# =========================
print("\n" + "=" * 55)
_API_KEY = os.getenv("GEMINI_API_KEY", "")
if not _API_KEY:
    raise RuntimeError("Missing GEMINI_API_KEY in .env file!")

print("[AutoDetect] Fetching available models from Gemini...")
_ALL_MODELS = list_all_gemini_models(_API_KEY)

EMBEDDING_BACKEND, EMBEDDING_DESC = detect_embedding_backend(_API_KEY, _ALL_MODELS)
LLM_MODEL_NAME,    LLM_DESC       = detect_llm(_API_KEY, _ALL_MODELS)

print(f"\n  Embedding : {EMBEDDING_DESC}")
print(f"  LLM       : {LLM_DESC}")
print("=" * 55 + "\n")


# =========================
# LLM FACTORY
# =========================
def get_llm():
    return ChatGoogleGenerativeAI(
        model=LLM_MODEL_NAME,
        temperature=0.2,
        google_api_key=_API_KEY,
    )


# =========================
# FILE READERS
# =========================
def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i+1}]\n{text}")
    return "\n\n".join(pages)

def read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def read_file(path: Path) -> str:
    if path.suffix == ".pdf":
        return read_pdf(path)
    if path.suffix == ".docx":
        return read_docx(path)
    return path.read_text(errors="ignore")


# =========================
# VECTOR DB
# =========================
def vector_db() -> Chroma:
    return Chroma(
        collection_name="rag_docs",
        persist_directory=str(VECTOR_DIR),
        embedding_function=EMBEDDING_BACKEND,
    )


# =========================
# IMPROVED RETRIEVAL
# =========================
def retrieve_chunks(db: Chroma, question: str, doc_id: str) -> list[Document]:
    """
    Fetch RETRIEVE_K chunks with scores, filter low-scoring ones,
    deduplicate overlapping chunks, return top FINAL_K.
    """
    results = db.similarity_search_with_relevance_scores(
        question,
        k=RETRIEVE_K,
        filter={"doc_id": doc_id},
    )

    # Filter by minimum score
    filtered = [(doc, score) for doc, score in results if score >= MIN_SCORE]

    if not filtered:
        # If nothing passes threshold, just return top results anyway
        filtered = results[:FINAL_K]

    # Deduplicate — remove chunks that are >80% the same as a higher-scored one
    seen: list[str] = []
    unique: list[Document] = []
    for doc, score in sorted(filtered, key=lambda x: x[1], reverse=True):
        content = doc.page_content.strip()
        is_duplicate = any(
            _overlap_ratio(content, s) > 0.8 for s in seen
        )
        if not is_duplicate:
            seen.append(content)
            unique.append(doc)
        if len(unique) >= FINAL_K:
            break

    return unique

def _overlap_ratio(a: str, b: str) -> float:
    """Simple word-overlap ratio between two strings."""
    words_a = set(a.lower().split())
    words_b = set(b.lower().split())
    if not words_a or not words_b:
        return 0.0
    return len(words_a & words_b) / min(len(words_a), len(words_b))


# =========================
# CONVERSATION HISTORY
# =========================
MAX_HISTORY_TURNS = 4  # keep last N exchanges per conversation

def get_history(conversation_id: str) -> list[dict]:
    if not conversation_id:
        return []
    history = load_history()
    return history.get(conversation_id, [])

def append_history(conversation_id: str, question: str, answer: str):
    if not conversation_id:
        return
    history = load_history()
    turns = history.get(conversation_id, [])
    turns.append({"role": "user",      "content": question})
    turns.append({"role": "assistant", "content": answer})
    # Keep only last MAX_HISTORY_TURNS exchanges
    turns = turns[-(MAX_HISTORY_TURNS * 2):]
    history[conversation_id] = turns
    save_history(history)

def format_history(turns: list[dict]) -> str:
    if not turns:
        return ""
    lines = []
    for t in turns:
        prefix = "User" if t["role"] == "user" else "Assistant"
        lines.append(f"{prefix}: {t['content']}")
    return "\n".join(lines)


# =========================
# APP
# =========================
app = FastAPI(title="RAG API — Improved")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def home():
    return {
        "status":            "RAG API running ✅",
        "embedding_backend": EMBEDDING_DESC,
        "llm_model":         LLM_DESC,
    }


@app.get("/api/models")
def show_models():
    all_models = list_all_gemini_models(_API_KEY)
    return {
        "active_embedding": EMBEDDING_DESC,
        "active_llm":       LLM_DESC,
        "all_embedding_models": [
            m["name"] for m in all_models
            if "embedContent" in m.get("supportedGenerationMethods", [])
        ],
        "all_llm_models": [
            m["name"] for m in all_models
            if "generateContent" in m.get("supportedGenerationMethods", [])
        ],
    }


# =========================
# UPLOAD
# =========================
@app.post("/api/documents/upload")
async def upload(file: UploadFile = File(...)):
    ensure_dirs()

    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(400, f"Unsupported type '{ext}'. Allowed: {SUPPORTED_EXTENSIONS}")

    doc_id    = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{doc_id}{ext}"

    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        text = read_file(save_path)
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(400, f"Could not read file: {e}")

    if not text.strip():
        save_path.unlink(missing_ok=True)
        raise HTTPException(400, "Document is empty or image-based (no extractable text)")

    # Split with semantic separators — tries to break at headings/paragraphs first
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n\n", "\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)

    docs = [
        Document(
            page_content=chunk,
            metadata={
                "doc_id":      doc_id,
                "source":      file.filename,
                "chunk_index": i,
                "total_chunks": len(chunks),
            },
        )
        for i, chunk in enumerate(chunks)
    ]

    try:
        db = vector_db()
        db.add_documents(docs)
    except Exception as e:
        save_path.unlink(missing_ok=True)
        raise HTTPException(500, f"Embedding failed: {e}")

    state = load_state()
    state.append({
        "id":            doc_id,
        "filename":      file.filename,
        "created_at":    now(),
        "chunks":        len(chunks),
        "embedded_with": EMBEDDING_DESC,
    })
    save_state(state)

    return {
        "doc_id":        doc_id,
        "filename":      file.filename,
        "chunks":        len(chunks),
        "embedded_with": EMBEDDING_DESC,
    }


# =========================
# DOCUMENTS
# =========================
@app.get("/api/documents")
def list_documents():
    return sorted(load_state(), key=lambda d: d.get("created_at", ""), reverse=True)


@app.delete("/api/documents/{document_id}")
def delete_document(document_id: str):
    state     = load_state()
    remaining = [d for d in state if d["id"] != document_id]
    if len(state) == len(remaining):
        raise HTTPException(404, "Document not found")
    save_state(remaining)

    try:
        db = vector_db()
        db.delete(where={"doc_id": document_id})
    except Exception:
        pass

    for f in UPLOAD_DIR.iterdir():
        if f.stem == document_id:
            f.unlink(missing_ok=True)
            break

    return {"status": "deleted", "id": document_id}


# =========================
# CHAT (RAG) — IMPROVED
# =========================
class ChatRequest(BaseModel):
    document_id:     str
    question:        str
    conversation_id: str | None = None


@app.post("/api/chat")
def chat(payload: ChatRequest):
    if not payload.question.strip():
        raise HTTPException(400, "Question cannot be empty")

    db   = vector_db()
    docs = retrieve_chunks(db, payload.question, payload.document_id)

    if not docs:
        return {"answer": "No relevant content found in this document for your question."}

    # Number the context chunks so the LLM knows they're separate sections
    context_parts = [
        f"[Section {i+1}]\n{doc.page_content}"
        for i, doc in enumerate(docs)
    ]
    context = "\n\n---\n\n".join(context_parts)

    # Load conversation history
    history       = get_history(payload.conversation_id)
    history_text  = format_history(history)

    # Build prompt — with or without history
    if history_text:
        prompt = ChatPromptTemplate.from_template(
            """You are a helpful assistant answering questions about a document.
Use ONLY the document sections below to answer. 
Combine information from multiple sections if needed.
If the answer truly isn't in the document, say so clearly.

--- DOCUMENT SECTIONS ---
{context}

--- CONVERSATION SO FAR ---
{history}

--- NEW QUESTION ---
{question}

Answer:"""
        )
        chain_input = {
            "context":  context,
            "history":  history_text,
            "question": payload.question,
        }
    else:
        prompt = ChatPromptTemplate.from_template(
            """You are a helpful assistant answering questions about a document.
Use ONLY the document sections below to answer.
Combine information from multiple sections if needed.
If the answer truly isn't in the document, say so clearly.

--- DOCUMENT SECTIONS ---
{context}

--- QUESTION ---
{question}

Answer:"""
        )
        chain_input = {
            "context":  context,
            "question": payload.question,
        }

    chain = prompt | get_llm()

    try:
        result = chain.invoke(chain_input)
    except Exception as e:
        raise HTTPException(500, f"LLM error: {e}")

    answer = result.content

    # Save to history
    append_history(payload.conversation_id, payload.question, answer)

    return {
        "answer":   answer,
        "sources":  [doc.metadata for doc in docs],
        "llm_used": LLM_DESC,
        "chunks_used": len(docs),
    }


# =========================
# CONVERSATION MANAGEMENT
# =========================
@app.get("/api/conversations/{conversation_id}")
def get_conversation(conversation_id: str):
    history = load_history()
    return {
        "conversation_id": conversation_id,
        "turns": history.get(conversation_id, []),
    }

@app.delete("/api/conversations/{conversation_id}")
def clear_conversation(conversation_id: str):
    history = load_history()
    if conversation_id in history:
        del history[conversation_id]
        save_history(history)
    return {"status": "cleared", "conversation_id": conversation_id}
